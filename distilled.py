#!/usr/bin/env python3
"""
Distilled — 每日 YouTube 精华文章自动生成系统
"""

import os, json, re, smtplib, logging, datetime, base64, urllib.request, math, time, html, io
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from pathlib import Path

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent / '.env')

from googleapiclient.discovery import build
from youtube_transcript_api import YouTubeTranscriptApi
from google import genai
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ebooklib import epub
from playwright.sync_api import sync_playwright

# ─── 配置 ──────────────────────────────────────────────────────────────────────

YOUTUBE_API_KEY  = os.getenv('YOUTUBE_API_KEY')
GEMINI_API_KEY   = os.getenv('GEMINI_API_KEY')
GMAIL_USER       = os.getenv('GMAIL_USER')
GMAIL_PASSWORD   = os.getenv('GMAIL_PASSWORD')
RECIPIENT_EMAIL  = os.getenv('RECIPIENT_EMAIL')
OUTPUT_DIR       = Path(os.getenv('OUTPUT_DIR', os.path.expanduser('~/Desktop/Distilled')))
PROCESSED_FILE   = OUTPUT_DIR / 'processed.json'

CHANNELS = [
    {'name': "Lenny's Podcast",  'handle': 'lennypodcast'},
    {'name': 'Andrej Karpathy',  'handle': 'AndrejKarpathy'},
    {'name': 'Lex Fridman',      'handle': 'lexfridman'},
    {'name': 'Y Combinator',     'handle': 'ycombinator'},
    {'name': 'Every',            'handle': 'everydotto'},
]

KEYWORDS = [
    'AI tools', 'vibe coding', 'vibe coder', 'startup', 'founder',
    'product management', 'artificial intelligence', 'LLM', 'AGI',
    'machine learning', 'product', 'growth',
]

# 全局关键词搜索查询（在整个 YouTube 范围内搜索）
SEARCH_QUERIES = [
    'vibe coding AI',
    'AI tools startup founder',
    'AI product management',
    'LLM AGI startup',
]

MIN_DURATION_MINUTES = 20
LOOKBACK_HOURS = 720  # 30 天

# 真正在做东西的人：加分
BUILDER_SIGNALS = [
    'how i built', 'i built', 'we built', 'i made', 'building',
    'open source', 'deep dive', 'architecture', 'engineering',
    'tutorial', 'code review', 'my project', 'developer', 'from scratch',
    'technical', 'implementation', 'behind the scenes', 'breakdown',
]

# 流量内容：扣分
INFLUENCER_SIGNALS = [
    'ranked', 'worst to best', 'best to worst', 'tier list',
    'nobody tells you', 'secret', 'you need to know', 'must know',
    'i made $', 'revenue:', 'day \\d+', 'days until', 'making $',
    'react to', 'reacting', 'watch me', 'i tried',
    'top \\d+', '\\d+ ways', '\\d+ mistakes', '\\d+ tips', '\\d+ things',
]

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s  %(levelname)-7s  %(message)s',
    datefmt='%H:%M:%S',
)
log = logging.getLogger('distilled')


# ─── YouTube ───────────────────────────────────────────────────────────────────

def get_channel_id(youtube, handle):
    """频道 handle → channel ID"""
    try:
        r = youtube.channels().list(part='id', forHandle=handle).execute()
        if r.get('items'):
            return r['items'][0]['id']
    except Exception:
        pass
    try:
        r = youtube.search().list(part='snippet', q=handle, type='channel', maxResults=1).execute()
        if r.get('items'):
            return r['items'][0]['snippet']['channelId']
    except Exception as e:
        log.warning(f'Cannot resolve channel "{handle}": {e}')
    return None


def parse_duration(s):
    """ISO 8601 时长 → 分钟数"""
    m = re.match(r'PT(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?', s or '')
    if not m:
        return 0
    h, mn, sc = int(m.group(1) or 0), int(m.group(2) or 0), int(m.group(3) or 0)
    return h * 60 + mn + sc / 60


def get_recent_videos(youtube, channel_id):
    since = (datetime.datetime.now(datetime.timezone.utc) - datetime.timedelta(hours=LOOKBACK_HOURS)
             ).strftime('%Y-%m-%dT%H:%M:%SZ')
    try:
        r = youtube.search().list(
            part='snippet', channelId=channel_id, type='video',
            order='date', publishedAfter=since, maxResults=5
        ).execute()
        ids = [item['id']['videoId'] for item in r.get('items', [])]
        if not ids:
            return []
        details = youtube.videos().list(part='contentDetails,snippet,statistics', id=','.join(ids)).execute()
        videos = []
        for item in details.get('items', []):
            stats = item.get('statistics', {})
            videos.append({
                'id':               item['id'],
                'title':            item['snippet']['title'],
                'published_at':     item['snippet']['publishedAt'],
                'duration_minutes': parse_duration(item['contentDetails']['duration']),
                'url':              f"https://www.youtube.com/watch?v={item['id']}",
                'like_count':       int(stats.get('likeCount',  0)),
                'view_count':       int(stats.get('viewCount',  0)),
            })
        return videos
    except Exception as e:
        log.error(f'Error fetching videos: {e}')
        return []


def score_video(video):
    title_lower = video['title'].lower()

    # 1. 互动量（主要依据）
    likes = video.get('like_count', 0)
    views = video.get('view_count', 0)
    engagement = math.log10(likes + 1) * 40 + math.log10(views + 1) * 10

    # 2. 关键词匹配
    keyword_score = sum(15 for kw in KEYWORDS if kw.lower() in title_lower)

    # 3. Builder 加分 / Influencer 扣分
    builder_score = sum(
        20 for sig in BUILDER_SIGNALS
        if re.search(sig, title_lower)
    )
    influencer_penalty = sum(
        25 for sig in INFLUENCER_SIGNALS
        if re.search(sig, title_lower)
    )

    # 4. 时效性（极小权重，仅做微调）
    recency = 0
    try:
        pub = datetime.datetime.fromisoformat(video['published_at'].replace('Z', '+00:00'))
        age_h = (datetime.datetime.now(datetime.timezone.utc) - pub).total_seconds() / 3600
        recency = max(0, 3 - age_h / 250)
    except Exception:
        pass

    return engagement + keyword_score + builder_score - influencer_penalty + recency


def search_global_videos(youtube):
    """在全站搜索关键词相关视频，返回候选列表"""
    since = (datetime.datetime.now(datetime.timezone.utc) - datetime.timedelta(hours=LOOKBACK_HOURS)
             ).strftime('%Y-%m-%dT%H:%M:%SZ')
    seen_ids = set()
    videos = []
    for query in SEARCH_QUERIES:
        try:
            r = youtube.search().list(
                part='snippet', q=query, type='video',
                order='relevance', publishedAfter=since,
                maxResults=5, relevanceLanguage='en',
            ).execute()
            ids = [item['id']['videoId'] for item in r.get('items', [])
                   if item['id']['videoId'] not in seen_ids]
            if not ids:
                continue
            details = youtube.videos().list(part='contentDetails,snippet,statistics', id=','.join(ids)).execute()
            for item in details.get('items', []):
                vid_id = item['id']
                seen_ids.add(vid_id)
                stats = item.get('statistics', {})
                videos.append({
                    'id':               vid_id,
                    'title':            item['snippet']['title'],
                    'channel_name':     item['snippet']['channelTitle'],
                    'published_at':     item['snippet']['publishedAt'],
                    'duration_minutes': parse_duration(item['contentDetails']['duration']),
                    'url':              f"https://www.youtube.com/watch?v={vid_id}",
                    'like_count':       int(stats.get('likeCount', 0)),
                    'view_count':       int(stats.get('viewCount', 0)),
                })
        except Exception as e:
            log.error(f'Global search error for "{query}": {e}')
    return videos


def select_best_video(youtube, processed_ids):
    seen_ids = set(processed_ids)
    candidates = []

    # 1. 指定频道的最新视频
    for ch in CHANNELS:
        channel_id = get_channel_id(youtube, ch['handle'])
        if not channel_id:
            continue
        log.info(f"Checking channel: {ch['name']}…")
        for v in get_recent_videos(youtube, channel_id):
            if v['id'] in seen_ids:
                continue
            if v['duration_minutes'] < MIN_DURATION_MINUTES:
                log.info(f"  ✗ Too short ({v['duration_minutes']:.0f}min): {v['title']}")
                continue
            seen_ids.add(v['id'])
            v['channel_name'] = ch['name']
            v['score'] = score_video(v)
            candidates.append(v)
            log.info(f"  ✓ {v['score']:.0f}pts  {v['duration_minutes']:.0f}min  {v['title']}")

    # 2. 全局关键词搜索
    log.info('Searching global keywords…')
    for v in search_global_videos(youtube):
        if v['id'] in seen_ids:
            continue
        if v['duration_minutes'] < MIN_DURATION_MINUTES:
            continue
        seen_ids.add(v['id'])
        v['score'] = score_video(v)
        candidates.append(v)
        log.info(f"  ✓ {v['score']:.0f}pts  {v['duration_minutes']:.0f}min  {v['title']}  [{v['channel_name']}]")

    if not candidates:
        return None
    return max(candidates, key=lambda x: x['score'])


# ─── 字幕提取 ──────────────────────────────────────────────────────────────────

def get_transcript(video_id):
    try:
        api = YouTubeTranscriptApi()
        t = api.fetch(video_id, languages=['en', 'en-US', 'en-GB'])
        return ' '.join(e.text for e in t)
    except Exception as e:
        raise RuntimeError(f'字幕提取失败：{e}')


# ─── 评论 ─────────────────────────────────────────────────────────────────────

def get_top_comments(youtube, video_id, max_results=40):
    """抓取视频热门评论，返回按点赞数排序的文本列表"""
    try:
        r = youtube.commentThreads().list(
            part='snippet', videoId=video_id,
            order='relevance', maxResults=max_results,
            textFormat='plainText',
        ).execute()
        comments = []
        for item in r.get('items', []):
            s = item['snippet']['topLevelComment']['snippet']
            comments.append({'text': s['textDisplay'], 'likes': s['likeCount']})
        comments.sort(key=lambda c: c['likes'], reverse=True)
        return [c['text'] for c in comments]
    except Exception as e:
        log.warning(f'无法获取评论：{e}')
        return []


# ─── 封面图 ────────────────────────────────────────────────────────────────────

def get_thumbnail(video_id):
    """下载 YouTube 视频封面，返回 (base64 data URL, 原始字节)"""
    for quality in ['maxresdefault', 'hqdefault', 'mqdefault']:
        url = f'https://img.youtube.com/vi/{video_id}/{quality}.jpg'
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            data = urllib.request.urlopen(req, timeout=10).read()
            b64 = base64.b64encode(data).decode()
            return f'data:image/jpeg;base64,{b64}', data
        except Exception:
            continue
    return None, None


# ─── Gemini ────────────────────────────────────────────────────────────────────

_gemini_client = None

def call_gemini(prompt, retries=2):
    global _gemini_client
    if _gemini_client is None:
        _gemini_client = genai.Client(api_key=GEMINI_API_KEY)
    for attempt in range(retries + 1):
        try:
            response = _gemini_client.models.generate_content(
                model='gemini-2.5-pro',
                contents=prompt,
            )
            return response.text.strip()
        except Exception as e:
            if attempt < retries:
                log.warning(f'Gemini 调用失败，30 秒后重试… ({e})')
                time.sleep(30)
            else:
                raise


ARTICLE_PROMPT = """\
你是一个有经验的中文杂志编辑，擅长从长篇播客中提炼精华，写出中文读者爱读的深度文章。

请根据下方字幕，生成一篇中文杂志文章，严格遵循以下要求：

【语言】纯正中文写作，不是翻译腔，前提是从一开始就会在中文杂志上发表。

【引言】开头段落自然传达：采访者/主讲人是谁及为何值得现在听 / 话题利害关系 / 让人"不读完会亏"的钩子。

【前提概念】识别字幕中反复出现的术语，用类比和日常例子解释，自然融入行文。目标读者：每天用 ChatGPT 但从未训练过模型的文科生。

【宝石级洞见】只提取：真正反直觉的观点 / 新颖视角 / 具体故事和例子 / 真实的自我披露 / 可落地的策略框架 / 影响日常工作和生活的观察。不提取：励志套话、广为人知的观点、表面评论、只有从业者才懂的技术细节。

【引用润色】去掉口语停顿词，纠正转录错误，合并重复表述，保留说话者语气。

【连接段落】宝石之间用简短桥接段，说明在讨论什么、为何重要、如何引出下一话题。

【文章格式】
- 完整叙事，无小标题，无章节分隔（允许用 --- 分节）
- 第一行：# 杂志风格标题
- 第二段：**斜体引言，概括文章内容和值得读的理由**
- 从头到尾自然流畅，不了解这个播客的人也能读懂
- 阅读时间约 5 分钟

【原文引用】遇到特别精彩、值得原文呈现的表达时，以 > 开头单独一行引用英文原文，紧接下一行以 > 译： 开头给出中文翻译，然后继续行文。每篇最多引用 5 处，只引用真正值得的句子。

只输出文章内容，不要任何额外说明。

---
视频标题：{title}
频道：{channel}
原视频：{url}

字幕：
{transcript}
"""


def generate_article(transcript, video_title, video_url, channel_name):
    prompt = ARTICLE_PROMPT.format(
        title=video_title, channel=channel_name, url=video_url,
        transcript=transcript[:80000]
    )
    return call_gemini(prompt)


def generate_filename(article_text):
    prompt = (
        "根据下面这篇文章，用一句简洁有力的中文概括文章核心内容（不超过18字），作为文件名。"
        "不要标点符号，不要引号，只输出那一句话。\n\n"
        f"文章（前600字）：\n{article_text[:600]}"
    )
    name = call_gemini(prompt)
    name = re.sub(r'[^\u4e00-\u9fffa-zA-Z0-9]', '', name)
    return name[:18]


EMAIL_PROMPT = """\
你是一位洞察人心的文案高手。根据下面的文章和视频评论，生成邮件素材。

严格按以下格式输出，每个区块用 ===区块名=== 标记：

===金句===
列出3条最有传播力的金句：
• [金句内容]

===故事===
列出3个最有共鸣的故事或场景，每个格式：
**[标题]**
[2-3句话描述]

===一句话===
列出3条25字以内的精华，适合做图片卡片：
• [内容]

===评论===
基于下方真实评论，总结3-4个讨论最热烈的话题，每个格式：
**[话题标题]**
[1-2句话概括讨论焦点]
❝ [原评论原文（英文）]（译：[中文翻译]）

===END===

文章内容：
{article}

---
视频热门评论（按点赞数排序）：
{comments}
"""


def generate_email_content(article_text, comments=None):
    comments_text = '\n'.join(f'• {c}' for c in (comments or [])[:30]) or '（暂无评论数据）'
    return call_gemini(EMAIL_PROMPT.format(article=article_text[:8000], comments=comments_text))


# ─── HTML 构建 ─────────────────────────────────────────────────────────────────

CSS = """
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;700&family=Noto+Sans+SC:wght@400;500&display=swap');
*,*::before,*::after{box-sizing:border-box}
body{margin:0;padding:0;background:#f7f4ef;color:#1c1c1c;
  font-family:'Noto Serif SC','Georgia','STSong',serif;font-size:17px;line-height:1.95;-webkit-font-smoothing:antialiased}
.page{max-width:720px;margin:0 auto;background:#fff;padding:64px 72px 96px;box-shadow:0 0 60px rgba(0,0,0,.07)}
.kicker{font-family:'Noto Sans SC',sans-serif;font-size:11px;font-weight:500;letter-spacing:.18em;text-transform:uppercase;color:#999;margin-bottom:18px}
h1.title{font-size:2em;font-weight:700;line-height:1.32;color:#111;margin:0 0 24px}
p.intro{font-size:1.06em;line-height:1.85;color:#444;border-left:3px solid #c8a96e;
  padding:3px 0 3px 20px;margin:0 0 32px;font-style:italic}
.meta{font-family:'Noto Sans SC',sans-serif;font-size:12px;color:#aaa;letter-spacing:.04em;
  margin-bottom:36px;padding-bottom:28px;border-bottom:1px solid #e8e4dd}
.meta a{color:#c8a96e;text-decoration:none;border-bottom:1px solid #c8a96e}
p{margin:0 0 1.25em;text-align:justify}
.sep{text-align:center;margin:32px 0;color:#c8a96e;font-size:13px;letter-spacing:.4em}
strong{font-weight:700;color:#111} em{font-style:italic;color:#333}
blockquote{border-left:3px solid #c8a96e;margin:1.4em 0;padding:6px 0 6px 20px;color:#555;font-size:0.95em;line-height:1.75}
blockquote p{margin:0;font-style:italic}
blockquote p.bq-trans{font-style:normal;font-size:0.88em;color:#999;margin-top:0.45em}
"""


def build_html(article_text, video_url, channel_name, date_str, thumbnail_data_url=None):
    title, intro, body_lines = '', '', []
    for raw in article_text.split('\n'):
        line = raw.strip()
        if line.startswith('# '):
            title = line[2:].strip()
        elif line.startswith('**') and line.endswith('**') and title != '':
            intro = line[2:-2].strip()
        elif line:
            body_lines.append(line)

    safe_title   = html.escape(title)
    safe_channel = html.escape(channel_name)

    def fmt(line):
        line = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', line)
        return re.sub(r'\*(.+?)\*', r'<em>\1</em>', line)

    body_html = ''
    i = 0
    while i < len(body_lines):
        line = body_lines[i]
        if line == '---':
            body_html += '<div class="sep">✦</div>\n'
            i += 1
        elif line.startswith('> '):
            parts = []
            while i < len(body_lines) and body_lines[i].startswith('> '):
                content = body_lines[i][2:]
                if content.startswith('译：'):
                    parts.append(f'<p class="bq-trans">{fmt(content[2:])}</p>')
                else:
                    parts.append(f'<p>{fmt(content)}</p>')
                i += 1
            body_html += f'<blockquote>{"".join(parts)}</blockquote>\n'
        else:
            body_html += f'<p>{fmt(line)}</p>\n'
            i += 1

    intro_html = f'<p class="intro">{fmt(intro)}</p>' if intro else ''

    cover_html = ''
    if thumbnail_data_url:
        cover_html = f'''<div class="cover-page">
  <img src="{thumbnail_data_url}" alt="封面"/>
  <div class="cover-meta">
    <div class="cover-channel">{safe_channel}</div>
    <div class="cover-title">{safe_title}</div>
    <div class="cover-date">{date_str}</div>
  </div>
</div>
<div class="page-break"></div>'''

    cover_css = """
.cover-page{width:100%;min-height:100vh;display:flex;flex-direction:column;
  justify-content:flex-end;background:#111;position:relative;overflow:hidden}
.cover-page img{position:absolute;top:0;left:0;width:100%;height:100%;
  object-fit:cover;opacity:0.75}
.cover-meta{position:relative;z-index:2;padding:40px 48px 48px;
  background:linear-gradient(transparent,rgba(0,0,0,0.85))}
.cover-channel{font-family:sans-serif;font-size:11px;letter-spacing:.2em;
  text-transform:uppercase;color:#c8a96e;margin-bottom:12px}
.cover-title{font-family:'Noto Serif SC','Georgia',serif;font-size:1.8em;
  font-weight:700;color:#fff;line-height:1.35;margin-bottom:12px}
.cover-date{font-family:sans-serif;font-size:11px;color:rgba(255,255,255,0.5)}
.page-break{page-break-after:always;break-after:page}
"""

    doc = f"""<!DOCTYPE html>
<html lang="zh"><head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>{safe_title}</title>
<style>{CSS}{cover_css}</style>
</head><body>
{cover_html}
<div class="page">
  <div class="kicker">{safe_channel} · Distilled</div>
  <h1 class="title">{safe_title}</h1>
  <div class="meta">{safe_channel} &nbsp;|&nbsp; {date_str} &nbsp;|&nbsp; <a href="{video_url}" target="_blank">观看原视频 ↗</a></div>
  {intro_html}
  <div class="body-text">{body_html}</div>
</div>
</body></html>"""
    return doc, title


# ─── 文件生成 ──────────────────────────────────────────────────────────────────

def make_pdf(html_content, output_path):
    # Strip Google Fonts import for offline rendering
    html_pdf = re.sub(r"@import url\('https://fonts\.googleapis\.com[^']*'\);", '', html_content)
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html_pdf, wait_until='domcontentloaded')
        page.wait_for_timeout(800)
        page.pdf(path=str(output_path), format='A4',
                 margin={'top':'20mm','bottom':'25mm','left':'20mm','right':'20mm'},
                 print_background=True)
        browser.close()


def make_epub(article_text, title, video_url, channel_name, date_str, thumbnail_bytes=None):
    _, intro, body_lines = '', '', []
    for raw in article_text.split('\n'):
        line = raw.strip()
        if line.startswith('# ') or not line:
            continue
        elif line.startswith('**') and line.endswith('**'):
            intro = line[2:-2].strip()
        else:
            body_lines.append(line)

    def fmt(line):
        line = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', line)
        return re.sub(r'\*(.+?)\*', r'<em>\1</em>', line)

    body_html = ''
    i = 0
    while i < len(body_lines):
        l = body_lines[i]
        if l == '---':
            body_html += '<hr/>\n'
            i += 1
        elif l.startswith('> '):
            parts = []
            while i < len(body_lines) and body_lines[i].startswith('> '):
                content = body_lines[i][2:]
                if content.startswith('译：'):
                    parts.append(f'<p class="bq-trans">{fmt(content[2:])}</p>')
                else:
                    parts.append(f'<p>{fmt(content)}</p>')
                i += 1
            body_html += f'<blockquote>{"".join(parts)}</blockquote>\n'
        else:
            body_html += f'<p>{fmt(l)}</p>\n'
            i += 1
    intro_html = f'<p class="intro">{fmt(intro)}</p>' if intro else ''

    xhtml = f"""<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="zh"><head>
<meta charset="utf-8"/><title>{title}</title>
<style>body{{font-family:Georgia,serif;font-size:1em;line-height:1.9;color:#1a1a1a}}
h1{{font-size:1.45em;line-height:1.35;margin-bottom:.2em}}
p.meta{{font-size:.8em;color:#999;font-style:italic;margin-bottom:2em}}
p.intro{{font-style:italic;color:#555;border-left:3px solid #c8a96e;padding-left:1em;margin:1.5em 0}}
p{{margin:0 0 1.1em;text-align:justify}}hr{{border:none;border-top:1px solid #ddd;margin:1.8em auto;width:40%}}
blockquote{{border-left:3px solid #c8a96e;margin:1.2em 0;padding:4px 0 4px 1em;color:#555}}
blockquote p{{margin:0;font-style:italic}}
blockquote p.bq-trans{{font-style:normal;font-size:.88em;color:#999;margin-top:.4em}}</style>
</head><body>
<h1>{title}</h1>
<p class="meta">{channel_name} · {date_str} · <a href="{video_url}">原视频</a></p>
{intro_html}{body_html}
</body></html>"""

    book = epub.EpubBook()
    book.set_identifier(f'distilled-{date_str}')
    book.set_title(title)
    book.set_language('zh')
    book.add_author(channel_name)

    # 封面图
    if thumbnail_bytes:
        book.set_cover('images/cover.jpg', thumbnail_bytes)

    ch = epub.EpubHtml(title=title, file_name='article.xhtml', lang='zh')
    ch.set_content(xhtml.encode('utf-8'))
    book.add_item(ch)
    book.toc = [epub.Link('article.xhtml', title, 'article')]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ['nav', ch]

    buf = io.BytesIO()
    epub.write_epub(buf, book, {})
    return buf.getvalue()


def make_word(word_content, base_filename, article_title, date_str):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3);  sec.right_margin  = Cm(3)

    def rfonts(run):
        rPr = run._r.get_or_add_rPr()
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:eastAsia'), '微软雅黑')
        rPr.insert(0, rf)

    def para(text, size=10.5, bold=False, italic=False,
             color=(40,40,40), indent=0, sb=0, sa=6, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
        run.font.color.rgb = RGBColor(*color)
        rfonts(run)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        if indent: p.paragraph_format.left_indent = Cm(indent)
        if align:  p.alignment = align
        return p

    # 标题
    para('内容洞察与传播素材库', 20, bold=True, color=(20,20,20),
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    para(article_title, 10, italic=True, color=(150,150,150),
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    para(f'来源：Distilled · {date_str}', 9, color=(180,180,180),
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=20)

    # 解析各区块
    sections_map = {}
    cur = None
    for line in word_content.split('\n'):
        line = line.strip()
        if re.match(r'^===\w+===$', line):
            cur = line.strip('=').strip()
            sections_map[cur] = []
        elif cur and line and line != 'END':
            sections_map[cur].append(line)

    display = [
        ('金句',   '一、最容易引爆传播的金句'),
        ('故事',   '二、最有共鸣的故事与场景'),
        ('洞见',   '三、最适合深度展开的洞见'),
        ('标题',   '四、可直接套用的标题模板'),
        ('一句话', '五、一句话精华（适合配图发布）'),
        ('评论',   '六、评论区最热话题'),
    ]

    for key, heading in display:
        if key not in sections_map:
            continue
        para(heading, 13, bold=True, color=(180,140,80), sb=18, sa=6)
        for line in sections_map[key]:
            if line.startswith('**') and line.endswith('**'):
                para(line.strip('*'), 11, bold=True, color=(50,50,50), sb=10, sa=3)
            elif line.startswith('→') or line.startswith('💡'):
                para(line, 9.5, italic=True, color=(150,110,50), indent=0.8, sa=8)
            elif line.startswith('❝'):
                para(line, 9.5, italic=True, color=(120,90,40), indent=1.2, sa=6)
            elif line.startswith('•'):
                para(line, 10.5, color=(40,40,40), indent=0.5, sa=3)
            elif re.match(r'^【.+】', line):
                para(line, 10, bold=True, color=(100,70,20), sb=8, sa=3)
            else:
                para(line, 10, color=(80,80,80), sa=3)

    para('本文档由 Distilled 自动生成 · 仅供内容创作参考',
         8.5, italic=True, color=(200,200,200),
         align=WD_ALIGN_PARAGRAPH.CENTER, sb=30)

    path = OUTPUT_DIR / f'{base_filename}.docx'
    doc.save(str(path))
    return path


# ─── 邮件 ──────────────────────────────────────────────────────────────────────

def word_content_to_html(word_content):
    """把 Gemini 生成的 ===区块=== 格式转成邮件 HTML"""
    sections_map = {}
    cur = None
    for line in word_content.split('\n'):
        line = line.strip()
        if re.match(r'^===\S+===$', line):
            cur = line.strip('=')
            sections_map[cur] = []
        elif cur and line and line != 'END':
            sections_map[cur].append(line)

    display = [
        ('金句',   '最有传播力的金句'),
        ('故事',   '最有共鸣的故事'),
        ('一句话', '一句话精华'),
        ('评论',   '评论区最热话题'),
    ]

    H3 = ('font-family:sans-serif;font-size:10px;font-weight:600;letter-spacing:.15em;'
          'text-transform:uppercase;color:#c8a96e;border-bottom:1px solid #ede8df;'
          'padding-bottom:8px;margin:32px 0 14px')

    out = ''
    for key, heading in display:
        lines = sections_map.get(key, [])
        if not lines:
            continue
        out += f'<h3 style="{H3}">{heading}</h3>\n'
        for line in lines:
            if line.startswith('**') and line.endswith('**'):
                text = line.strip('*')
                out += f'<p style="font-weight:700;color:#111;margin:14px 0 3px;font-size:14px">{text}</p>\n'
            elif line.startswith('→'):
                out += (f'<p style="color:#c8a96e;font-size:12px;margin:2px 0 10px;'
                        f'padding-left:14px;font-style:italic">{line}</p>\n')
            elif line.startswith('💡'):
                out += (f'<p style="color:#a07840;font-size:13px;margin:10px 0 4px;'
                        f'padding-left:14px;font-style:italic">{line}</p>\n')
            elif line.startswith('•'):
                out += (f'<p style="margin:6px 0;padding-left:14px;color:#333;'
                        f'font-size:14px;line-height:1.75">{line}</p>\n')
            elif re.match(r'^【.+】', line):
                out += (f'<p style="font-weight:600;color:#7a5220;margin:12px 0 5px;'
                        f'font-size:12px;letter-spacing:.03em">{line}</p>\n')
            elif line.startswith('❝'):
                out += (f'<blockquote style="border-left:3px solid #e8d5b0;margin:6px 0 12px 14px;'
                        f'padding:6px 14px;color:#777;font-style:italic;font-size:12px;'
                        f'background:#faf7f2;line-height:1.7">{line}</blockquote>\n')
            elif line:
                out += f'<p style="margin:4px 0;color:#555;font-size:13px;line-height:1.7">{line}</p>\n'
    return out


def _fmt_num(n):
    if n >= 1_000_000: return f'{n/1_000_000:.1f}M'
    if n >= 1_000:     return f'{n/1_000:.1f}K'
    return str(n)


def send_email(epub_bytes, article_title, video_url,
               channel_name, date_str, base_filename, word_content='',
               like_count=0, view_count=0):
    msg = MIMEMultipart()
    msg['From']    = GMAIL_USER
    msg['To']      = RECIPIENT_EMAIL
    msg['Subject'] = f'📖 Distilled · {date_str} · {article_title}'

    word_html    = word_content_to_html(word_content) if word_content else ''
    safe_title   = html.escape(article_title)
    safe_channel = html.escape(channel_name)

    body = f"""<div style="font-family:'Noto Serif SC',Georgia,serif;max-width:660px;margin:0 auto;background:#f7f4ef">
  <!-- 头部 -->
  <div style="background:#fff;padding:36px 40px 28px">
    <p style="font-family:sans-serif;font-size:10px;color:#bbb;letter-spacing:.15em;text-transform:uppercase;margin:0 0 14px">DISTILLED · {date_str}</p>
    <h2 style="font-size:1.4em;font-weight:700;line-height:1.35;margin:0 0 10px;color:#111">{safe_title}</h2>
    <p style="color:#999;font-family:sans-serif;font-size:12px;margin:0 0 6px">{safe_channel}</p>
    <p style="font-family:sans-serif;font-size:11px;color:#bbb;margin:0 0 18px">👍 {_fmt_num(like_count)} &nbsp;·&nbsp; 👁 {_fmt_num(view_count)}</p>
    <a href="{video_url}" style="font-family:sans-serif;color:#c8a96e;font-size:12px;text-decoration:none;border-bottom:1px solid #c8a96e;padding-bottom:1px">▶ 观看原视频</a>
  </div>
  <!-- 传播素材 -->
  <div style="background:#fff;margin-top:2px;padding:8px 40px 36px">
    {word_html}
  </div>
  <!-- 附件说明 -->
  <div style="padding:18px 40px;text-align:center">
    <p style="font-family:sans-serif;font-size:11px;color:#bbb;margin:0">
      附件：📗 <b>{base_filename}.epub</b> — 导入 Apple Books / Kindle 阅读
    </p>
  </div>
</div>"""
    msg.attach(MIMEText(body, 'html'))

    # EPUB 附件
    part = MIMEBase('application', 'epub+zip')
    part.set_payload(epub_bytes)
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment', filename=f'{base_filename}.epub')
    msg.attach(part)

    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as s:
        s.login(GMAIL_USER, GMAIL_PASSWORD)
        s.send_message(msg)
    log.info(f'Email → {RECIPIENT_EMAIL}')


# ─── 已处理记录 ────────────────────────────────────────────────────────────────

def load_processed():
    if PROCESSED_FILE.exists():
        return set(json.loads(PROCESSED_FILE.read_text()))
    return set()

def save_processed(ids):
    PROCESSED_FILE.write_text(json.dumps(list(ids), indent=2, ensure_ascii=False))


# ─── 主流程 ────────────────────────────────────────────────────────────────────

def main():
    log.info('=' * 50)
    log.info('Distilled 启动')
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    processed = load_processed()
    youtube   = build('youtube', 'v3', developerKey=YOUTUBE_API_KEY)
    date_str  = datetime.date.today().strftime('%Y-%m-%d')

    # 1. 选视频
    video = select_best_video(youtube, processed)
    if not video:
        log.info('今天没有符合条件的新视频，结束。')
        return

    log.info(f'选中：{video["title"]}  ({video["duration_minutes"]:.0f}min)')

    # 2. 字幕
    log.info('提取字幕…')
    try:
        transcript = get_transcript(video['id'])
    except RuntimeError as e:
        log.error(str(e))
        log.info('无法获取字幕，今日跳过。')
        return
    log.info(f'字幕长度：{len(transcript)} 字符')

    # 3. 文章
    log.info('Gemini 生成文章…')
    article_text = generate_article(transcript, video['title'], video['url'], video['channel_name'])

    # 4. 文件名
    log.info('生成中文文件名…')
    base_filename = f'{generate_filename(article_text)}_{date_str}'
    log.info(f'文件名：{base_filename}')

    # 5. 封面图
    log.info('下载视频封面…')
    thumbnail_data_url, thumbnail_bytes = get_thumbnail(video['id'])
    if thumbnail_data_url:
        log.info('封面图下载成功')
    else:
        log.warning('封面图下载失败，跳过')

    # 6. HTML → PDF（本地）
    html_content, article_title = build_html(article_text, video['url'], video['channel_name'], date_str, thumbnail_data_url)
    pdf_path = OUTPUT_DIR / f'{base_filename}.pdf'
    log.info('生成 PDF…')
    make_pdf(html_content, pdf_path)
    log.info(f'PDF → {pdf_path}')

    # 7. EPUB（邮件用）
    log.info('生成 EPUB…')
    epub_bytes = make_epub(article_text, article_title, video['url'], video['channel_name'], date_str, thumbnail_bytes)

    # 8. 评论 + 邮件素材
    log.info('抓取视频评论…')
    comments = get_top_comments(youtube, video['id'])
    log.info(f'获取到 {len(comments)} 条评论')
    log.info('Gemini 生成邮件素材…')
    email_content = generate_email_content(article_text, comments)

    # 9. 发邮件
    log.info('发送邮件…')
    send_email(epub_bytes, article_title, video['url'],
               video['channel_name'], date_str, base_filename, email_content,
               like_count=video.get('like_count', 0),
               view_count=video.get('view_count', 0))

    # 10. 标记已处理
    processed.add(video['id'])
    save_processed(processed)

    log.info(f'完成！本地文件：{OUTPUT_DIR}')
    log.info('=' * 50)


if __name__ == '__main__':
    main()
